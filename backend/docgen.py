"""
Xera AI — Document Generator
Professional PDF, DOCX, XLSX creation with rich layout engine.
"""

import html
import re
import time
import uuid
from pathlib import Path

DOCS_DIR = Path("/tmp/xera-docs")
_FILE_TTL_SECONDS = 86400

_MAX_WORD = 52   # break tokens longer than this in regular text
_MAX_CODE = 72   # wrap code lines at this


def _ensure_dir():
    DOCS_DIR.mkdir(exist_ok=True)


def cleanup_old_docs():
    if not DOCS_DIR.exists():
        return
    cutoff = time.time() - _FILE_TTL_SECONDS
    for f in DOCS_DIR.iterdir():
        if f.is_file() and f.stat().st_mtime < cutoff:
            try:
                f.unlink()
            except OSError:
                pass


# ─── Color System ─────────────────────────────────────────────────────────────

_THEMES = {
    "purple":  {"header": "#7c3aed", "h1": "#5b21b6"},
    "blue":    {"header": "#0369a1", "h1": "#0c4a6e"},
    "green":   {"header": "#065f46", "h1": "#064e3b"},
    "dark":    {"header": "#1e293b", "h1": "#334155"},
    "warm":    {"header": "#b45309", "h1": "#92400e"},
    "minimal": {"header": "#111827", "h1": "#111827"},
}

_COLOR_NAMES = {
    "rot": "#dc2626", "red": "#dc2626",
    "gelb": "#ca8a04", "yellow": "#ca8a04",
    "orange": "#ea580c",
    "pink": "#db2777", "rosa": "#db2777",
    "grau": "#4b5563", "grey": "#4b5563", "gray": "#4b5563",
    "türkis": "#0891b2", "teal": "#0d9488", "cyan": "#0891b2",
    "indigo": "#4338ca",
    "lila": "#7c3aed", "violet": "#7c3aed", "purple": "#7c3aed",
    "blau": "#0369a1", "blue": "#0369a1",
    "grün": "#059669", "gruen": "#059669", "green": "#059669",
    "dunkel": "#1e293b", "dark": "#1e293b",
    "amber": "#d97706", "warm": "#b45309",
    "minimal": "#111827", "schwarz": "#111827", "black": "#111827",
    "gold": "#b45309", "silber": "#475569", "silver": "#475569",
}


def _theme_from_hex(hex_color: str) -> dict:
    h = hex_color.lstrip("#")
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)

    def mix(c, w, ratio):
        return int(c * ratio + w * (1 - ratio))

    h1 = f"#{mix(r,0,.72):02x}{mix(g,0,.72):02x}{mix(b,0,.72):02x}"
    alt = f"#{mix(r,255,.07):02x}{mix(g,255,.07):02x}{mix(b,255,.07):02x}"
    brd = f"#{mix(r,255,.22):02x}{mix(g,255,.22):02x}{mix(b,255,.22):02x}"
    light = f"#{mix(r,255,.12):02x}{mix(g,255,.12):02x}{mix(b,255,.12):02x}"
    bright = (r * 299 + g * 587 + b * 114) / 1000
    htxt = "#ffffff" if bright < 145 else alt

    return {
        "header": hex_color, "header_text": htxt,
        "h1": h1, "ink": "#1e293b",
        "tbl_head": hex_color, "tbl_alt": alt,
        "code_bg": alt, "code_border": brd,
        "hr": brd, "accent": hex_color,
        "accent_light": light,
    }


def _darken(hex_color: str, ratio: float = 0.75) -> str:
    h = hex_color.lstrip("#")
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    return f"#{int(r*ratio):02x}{int(g*ratio):02x}{int(b*ratio):02x}"


def _resolve_theme(spec: str) -> dict:
    spec = (spec or "purple").lower().strip()
    if spec in _THEMES:
        return _theme_from_hex(_THEMES[spec]["header"])
    if spec in _COLOR_NAMES:
        return _theme_from_hex(_COLOR_NAMES[spec])
    clean = spec.lstrip("#")
    if re.fullmatch(r'[0-9a-f]{6}', clean):
        return _theme_from_hex(f"#{clean}")
    if re.fullmatch(r'[0-9a-f]{3}', clean):
        return _theme_from_hex(f"#{''.join(c*2 for c in clean)}")
    return _theme_from_hex("#7c3aed")


_LAYOUTS = {
    "compact":  {"title": 20, "h1": 13, "h2": 11, "h3": 10, "body": 9,  "code": 7.5, "lm": 0.88, "sm": 0.72},
    "normal":   {"title": 26, "h1": 15, "h2": 12, "h3": 11, "body": 10, "code": 8.5, "lm": 1.0,  "sm": 1.0},
    "spacious": {"title": 30, "h1": 18, "h2": 14, "h3": 12, "body": 12, "code": 10,  "lm": 1.18, "sm": 1.35},
}

_LAYOUT_ALIASES = {
    "kompakt": "compact", "eng": "compact", "small": "compact",
    "gross": "spacious", "groß": "spacious", "large": "spacious", "luftig": "spacious",
}


def _resolve_layout(spec: str) -> dict:
    key = _LAYOUT_ALIASES.get((spec or "normal").lower().strip(), (spec or "normal").lower().strip())
    return _LAYOUTS.get(key, _LAYOUTS["normal"])


_CODE_EXTENSIONS = {
    "py": ".py", "python": ".py", "sh": ".sh", "bash": ".sh",
    "js": ".js", "javascript": ".js", "ts": ".ts", "typescript": ".ts",
    "html": ".html", "htm": ".html", "css": ".css", "json": ".json",
    "yaml": ".yaml", "yml": ".yaml", "md": ".md", "markdown": ".md",
    "txt": ".txt", "text": ".txt", "ps1": ".ps1", "powershell": ".ps1",
    "sql": ".sql", "go": ".go", "rs": ".rs", "rust": ".rs",
    "java": ".java", "c": ".c", "cpp": ".cpp",
    "toml": ".toml", "ini": ".ini", "conf": ".conf",
}


# ─── Content sanitization ──────────────────────────────────────────────────────
# Defense-in-depth: the LLM is instructed to keep create_document(content=...)
# limited to pure document markdown, but it sometimes ignores that and appends
# its own chat-style follow-up text or a manual markdown-link "table of contents"
# to the end of the content string. Since this function is the single choke
# point every caller (real tool call, auto-doc fallback, future callers) goes
# through before rendering, strip both patterns here so the bug can never leak
# into a generated file regardless of which code path produced the content.

_CHAT_LEAK_STARTS = (
    "✅", "✓", "⭐", "■", "fertig!", "fertig ", "was soll ich", "ich habe ", "möchtest du",
    "anpassungsmöglich", "erstellt am:", "erstellt mit xera",
    "was willst du", "was möchtest du", "soll ich noch",
    "what would you", "would you like", "feel free to", "done!", "all set",
)

_MANUAL_TOC_RE = re.compile(
    r'\n#{1,3}[ \t]*(?:Inhaltsverzeichnis|Table of [Cc]ontents)[ \t]*\n'
    r'(?:[ \t]*(?:\d+\.|[-*])[ \t]*\[[^\]]+\]\([^)]*\)[ \t]*\n?)+',
    re.IGNORECASE,
)


def _strip_manual_toc(content: str) -> str:
    """Remove a model-written markdown-link TOC block (we render real TOCs via toc=True)."""
    return _MANUAL_TOC_RE.sub('\n', content)


def _strip_chat_leak(content: str) -> str:
    """Remove trailing chat/follow-up lines (e.g. '✅ Fertig! Was soll ich anpassen?')
    that the model sometimes appends inside the content argument instead of the chat reply."""
    lines = content.split('\n')
    end = len(lines)
    for i in range(len(lines) - 1, -1, -1):
        stripped = lines[i].strip().lower()
        if not stripped:
            continue
        if (any(stripped.startswith(p) for p in _CHAT_LEAK_STARTS)
                or stripped.startswith('→') or stripped.startswith('-> ')
                or stripped.startswith('— ') or stripped.startswith('-- ')):
            end = i
        else:
            break
    while end > 0 and lines[end - 1].strip() in ('---', ''):
        end -= 1
    return '\n'.join(lines[:end]).rstrip()


def _sanitize_content(content: str) -> str:
    return _strip_chat_leak(_strip_manual_toc(content))


# ─── Public API ───────────────────────────────────────────────────────────────

def create_document(
    doc_type: str, content: str, filename: str,
    theme: str = "purple", layout: str = "normal",
    show_header: bool = True, cover_page: bool = False,
    toc: bool = False,
    header_left: str = "Xera AI",
    header_right: str = "xera-app.com",
) -> str:
    _ensure_dir()
    uid = uuid.uuid4().hex[:8]
    safe = "".join(c for c in filename if c.isalnum() or c in "._- ")[:50].strip() or "dokument"
    base = DOCS_DIR / f"{uid}_{safe}"
    t = _resolve_theme(theme)
    lay = _resolve_layout(layout)

    if doc_type in ("pdf", "docx"):
        content = _sanitize_content(content)

    if doc_type == "pdf":
        path = base.with_suffix(".pdf")
        _create_pdf(content, path, t, lay, show_header=show_header, cover_page=cover_page,
                    toc=toc, header_left=header_left, header_right=header_right)
        return f"/api/download/{path.name}"
    elif doc_type == "docx":
        path = base.with_suffix(".docx")
        _create_docx(content, path, t)
        return f"/api/download/{path.name}"
    elif doc_type == "xlsx":
        path = base.with_suffix(".xlsx")
        _create_xlsx(content, path, t)
        return f"/api/download/{path.name}"
    elif doc_type.lower() in _CODE_EXTENSIONS:
        path = base.with_suffix(_CODE_EXTENSIONS[doc_type.lower()])
        path.write_text(content, encoding="utf-8")
        return f"/api/download/{path.name}"
    else:
        path = base.with_suffix(".txt")
        path.write_text(content, encoding="utf-8")
        return f"/api/download/{path.name}"


# ─── Text Processing ──────────────────────────────────────────────────────────

def _break_long_str(s: str, n: int) -> str:
    if len(s) <= n:
        return s
    return " ".join(s[i:i+n] for i in range(0, len(s), n))


def _truncate_words(s: str, n: int) -> str:
    """Truncate at a word boundary (never mid-word) and add an ellipsis if cut."""
    if len(s) <= n:
        return s
    cut = s[:n].rsplit(" ", 1)[0].rstrip(",.;:–-")
    if not cut:
        cut = s[:n]
    return cut + "…"


def _preprocess(line: str) -> str:
    segs = re.split(r'(`[^`]*`)', line)
    out = []
    for seg in segs:
        if seg.startswith("`") and seg.endswith("`") and len(seg) > 2:
            out.append(f"`{_break_long_str(seg[1:-1], _MAX_WORD)}`")
        else:
            out.append(" ".join(_break_long_str(w, _MAX_WORD) for w in seg.split(" ")))
    return "".join(out)


def _md(text: str) -> str:
    t = html.escape(text)
    t = re.sub(r'\*\*\*(.+?)\*\*\*', r'<b><i>\1</i></b>', t)
    t = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', t)
    t = re.sub(r'\*([^*]+?)\*', r'<i>\1</i>', t)
    t = re.sub(r'`([^`]+?)`', lambda m: f'<font name="Courier" size="9">{m.group(1)}</font>', t)
    return t


def _p(text: str, style):
    from reportlab.platypus import Paragraph
    try:
        return Paragraph(_md(_preprocess(text)), style)
    except Exception:
        from reportlab.platypus import Paragraph as _P2
        return _P2(html.escape(text[:400]), style)


def _wrap_code(text: str) -> str:
    out = []
    for line in text.split("\n"):
        if len(line) <= _MAX_CODE:
            out.append(line)
        else:
            while len(line) > _MAX_CODE:
                out.append(line[:_MAX_CODE])
                line = "  " + line[_MAX_CODE:]
            out.append(line)
    return "\n".join(out)


# ─── Callout Boxes ────────────────────────────────────────────────────────────

# Callout type → (strip_color, background_color, label)
# Fixed colors per type — independent of document theme so they always look distinct.
# NOTE: Helvetica (WinAnsi/cp1252 encoding) cannot render ★ℹ⚠✓✕ — they show as
# missing-glyph boxes ("■") in real PDF viewers. Use text-only labels instead;
# the colored strip + background already carries the visual distinction.
_CALLOUT_CFG = {
    "key":     ("#6d28d9", "#ede9fe", "KERNAUSSAGE"),   # violet
    "info":    ("#0369a1", "#dbeafe", "INFO"),          # blue
    "warning": ("#b45309", "#fef3c7", "WARNUNG"),       # amber
    "tip":     ("#047857", "#d1fae5", "TIPP"),          # green
    "note":    ("#4b5563", "#f3f4f6", "HINWEIS"),       # gray
    "error":   ("#b91c1c", "#fee2e2", "FEHLER"),        # red
}


def _make_callout(content_lines, ctype, title, accent_hex, alt_hex, doc_width, fs_body, lm):
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib.styles import ParagraphStyle

    # Use fixed callout colors (not theme-dependent) for visual distinctness
    bc_hex, bg_hex, auto_label = _CALLOUT_CFG.get(
        ctype or "", (accent_hex, "#f0f4ff", "INFO")
    )
    c_bc  = colors.HexColor(bc_hex)
    c_bg  = colors.HexColor(bg_hex)
    c_txt = colors.HexColor("#1e293b")

    label_s = ParagraphStyle("CALL", fontSize=fs_body - 1, fontName="Helvetica-Bold",
                              textColor=c_bc, leading=(fs_body - 1) * 1.3, spaceAfter=3)
    body_s  = ParagraphStyle("CALB", fontSize=fs_body, fontName="Helvetica",
                              textColor=c_txt, leading=fs_body * lm * 1.55, spaceAfter=2)

    STRIP = 8       # wider strip for better visibility
    content_w = doc_width - STRIP

    paras = []
    # Show label+title as header, or just label if no title
    header_text = f"{auto_label}  —  {title}" if title else auto_label
    paras.append(("label", header_text))
    for ln in content_lines:
        ln = ln.strip()
        if ln:
            paras.append(("body", ln))

    if len(paras) <= 1 and not content_lines:
        return None

    rows = []
    for kind, text in paras:
        s = label_s if kind == "label" else body_s
        rows.append(["", _p(text, s)])

    tbl = Table(rows, colWidths=[STRIP, content_w])
    ts = [
        ("SPAN",          (0, 0), (0, -1)),
        ("BACKGROUND",    (0, 0), (0, -1), c_bc),
        ("BACKGROUND",    (1, 0), (1, -1), c_bg),
        ("BOX",           (0, 0), (-1, -1), 1,   c_bc),
        ("LEFTPADDING",   (0, 0), (0, -1), 0),
        ("RIGHTPADDING",  (0, 0), (0, -1), 0),
        ("LEFTPADDING",   (1, 0), (1, -1), 12),
        ("RIGHTPADDING",  (1, 0), (1, -1), 10),
        ("TOPPADDING",    (0, 0), (-1,  0), 10),
        ("BOTTOMPADDING", (0, -1), (-1, -1), 10),
        ("TOPPADDING",    (1, 1), (1, -1), 3),
        ("BOTTOMPADDING", (1, 0), (1, -2), 3),
        ("VALIGN",        (0, 0), (-1, -1), "TOP"),
    ]
    tbl.setStyle(TableStyle(ts))
    return tbl


# ─── Section Header (## level) — full-color band ─────────────────────────────

def _make_section_header(text, h1_hex, doc_width, fs_h1, lm):
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib.styles import ParagraphStyle

    c_h1   = colors.HexColor(h1_hex)
    c_text = colors.white

    s = ParagraphStyle("SH", fontSize=fs_h1, fontName="Helvetica-Bold",
                       textColor=c_text, leading=fs_h1 * lm * 1.3)
    # Full-width colored band: 4px darker left accent + full-width h1 background
    row = [["", _p(text, s)]]
    tbl = Table(row, colWidths=[4, doc_width - 4])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(_darken(h1_hex, 0.7))),
        ("BACKGROUND", (1, 0), (1, -1), c_h1),  # full colored band, white text
        ("LEFTPADDING",  (0, 0), (0, -1), 0),
        ("RIGHTPADDING", (0, 0), (0, -1), 0),
        ("LEFTPADDING",  (1, 0), (1, -1), 14),
        ("RIGHTPADDING", (1, 0), (1, -1), 10),
        ("TOPPADDING",   (0, 0), (-1, -1), 9),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 9),
        ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
    ]))
    return tbl


# ─── PDF Creator ──────────────────────────────────────────────────────────────

def _create_pdf(content, path, theme, lay, show_header=True, cover_page=False,
                toc=False, header_left="Xera AI", header_right="xera-app.com"):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm, mm
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY, TA_CENTER
    from reportlab.platypus import (
        SimpleDocTemplate, BaseDocTemplate, PageTemplate, Frame,
        Paragraph, Spacer, HRFlowable,
        Table, TableStyle, Preformatted, PageBreak,
    )
    from reportlab.platypus.tableofcontents import TableOfContents

    W, H = A4
    PAD = 10

    lm = lay["lm"]
    sm = lay["sm"]
    fs = lay

    c_header  = colors.HexColor(theme["header"])
    c_htxt    = colors.HexColor(theme["header_text"])
    c_h1      = colors.HexColor(theme["h1"])
    c_ink     = colors.HexColor(theme["ink"])
    c_tblhd   = colors.HexColor(theme["tbl_head"])
    c_tblalt  = colors.HexColor(theme["tbl_alt"])
    c_codebg  = colors.HexColor(theme["code_bg"])
    c_codebrd = colors.HexColor(theme["code_border"])
    c_hr      = colors.HexColor(theme["hr"])
    c_body    = colors.HexColor("#374151")
    c_muted   = colors.HexColor("#6b7280")
    white     = colors.white

    # — Styles — splitLongWords prevents ANY text from going past the margin —
    _SW = {"splitLongWords": 1, "wordWrap": "LTR"}

    title_s = ParagraphStyle("T", fontSize=fs["title"], fontName="Helvetica-Bold",
                             textColor=c_ink, leading=fs["title"]*lm*1.15, spaceAfter=3*sm, **_SW)
    h1_s    = ParagraphStyle("H1", fontSize=fs["h1"], fontName="Helvetica-Bold",
                             textColor=c_h1, leading=fs["h1"]*lm*1.3, **_SW)
    h2_s    = ParagraphStyle("H2", fontSize=fs["h2"], fontName="Helvetica-Bold",
                             textColor=c_ink, spaceBefore=12*sm, spaceAfter=3*sm,
                             leading=fs["h2"]*lm*1.3, **_SW)
    h3_s    = ParagraphStyle("H3", fontSize=fs["h3"], fontName="Helvetica-Bold",
                             textColor=c_body, spaceBefore=8*sm, spaceAfter=2*sm,
                             leading=fs["h3"]*lm*1.3, **_SW)
    body_s  = ParagraphStyle("B", fontSize=fs["body"], fontName="Helvetica",
                             textColor=c_body, leading=fs["body"]*lm*1.62,
                             spaceAfter=4*sm, alignment=TA_JUSTIFY, **_SW)
    blt_s   = ParagraphStyle("BL", fontSize=fs["body"], fontName="Helvetica",
                             textColor=c_body, leading=fs["body"]*lm*1.5,
                             leftIndent=14, spaceAfter=2*sm, **_SW)
    blt2_s  = ParagraphStyle("BL2", fontSize=fs["body"]-0.5, fontName="Helvetica",
                             textColor=c_muted, leading=fs["body"]*lm*1.5,
                             leftIndent=28, spaceAfter=2*sm, **_SW)
    num_s   = ParagraphStyle("NL", fontSize=fs["body"], fontName="Helvetica",
                             textColor=c_body, leading=fs["body"]*lm*1.5,
                             leftIndent=18, firstLineIndent=-14, spaceAfter=2*sm, **_SW)
    code_s  = ParagraphStyle("C", fontSize=fs["code"], fontName="Courier",
                             textColor=colors.HexColor("#1f2937"),
                             leading=fs["code"]*lm*1.42, spaceBefore=0, spaceAfter=0)

    lines = content.split("\n")
    doc_title_raw = next((ln[2:].strip() for ln in lines if ln.startswith("# ")), "")
    doc_title = _truncate_words(doc_title_raw, 70)

    # — Page drawing callbacks —
    def _draw_header(canvas, doc, first=False):
        if not show_header:
            return
        canvas.saveState()
        hh = 22*mm if first else 8*mm
        canvas.setFillColor(c_header)
        canvas.rect(0, H - hh, W, hh, fill=1, stroke=0)
        if first:
            canvas.setFont("Helvetica-Bold", 11)
            canvas.setFillColor(white)
            canvas.drawString(2.5*cm, H - 14*mm, _truncate_words(header_left, 30))
            canvas.setFont("Helvetica", 8)
            canvas.setFillColor(c_htxt)
            if doc_title:
                canvas.drawCentredString(W/2, H - 14*mm, _truncate_words(doc_title, 55))
            canvas.drawRightString(W - 2.5*cm, H - 14*mm, _truncate_words(header_right, 30))
        else:
            canvas.setFont("Helvetica-Bold", 8)
            canvas.setFillColor(white)
            canvas.drawString(2.5*cm, H - 5.5*mm, _truncate_words(header_left, 30))
            if doc_title:
                canvas.setFillColor(c_htxt)
                canvas.drawRightString(W - 2.5*cm, H - 5.5*mm, _truncate_words(doc_title, 50))
        # Footer
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(c_muted)
        footer_left = f"Erstellt mit {header_left}" if header_left != "Xera AI" else "Erstellt mit Xera AI · xera-app.com"
        canvas.drawString(2.5*cm, 11*mm, footer_left)
        canvas.drawRightString(W - 2.5*cm, 11*mm, f"Seite {doc.page}")
        canvas.setStrokeColor(colors.HexColor("#e5e7eb"))
        canvas.setLineWidth(0.4)
        canvas.line(2.5*cm, 15*mm, W - 2.5*cm, 15*mm)
        canvas.restoreState()

    class _TocAwareDocTemplate(BaseDocTemplate):
        def afterFlowable(self, flowable):
            entry = getattr(flowable, "_toc_entry", None)
            if entry:
                level, text = entry
                self.notify("TOCEntry", (level, text, self.page))

    margins = dict(
        topMargin=(3.0 if show_header else 1.5)*cm,
        bottomMargin=2.5*cm,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
    )
    # Proper PDF Info dictionary (Title/Author/Producer) — avoids "Untitled"/blank
    # metadata that some PDF readers flag or warn about on download, and makes
    # the file look professionally generated rather than auto-spat-out.
    meta = dict(
        title=doc_title or path.stem.split("_", 1)[-1].replace("_", " ").strip().title() or "Dokument",
        author=header_left if header_left != "Xera AI" else "Xera AI",
        subject="Erstellt mit Xera AI",
        creator="Xera AI (xera-app.com)",
        producer="Xera AI Document Engine",
        displayDocTitle=True,
    )

    if toc:
        doc = _TocAwareDocTemplate(str(path), pagesize=A4, **margins, **meta)
        frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")

        def _on_page(canvas, d):
            _draw_header(canvas, d, first=(d.page == 1))

        doc.addPageTemplates([PageTemplate(id="all", frames=[frame], onPage=_on_page)])
    else:
        doc = SimpleDocTemplate(str(path), pagesize=A4, **margins, **meta)

    story = []

    # — Optional cover page —
    if cover_page and doc_title:
        cover_t_s = ParagraphStyle("CT", fontSize=fs["title"]+8, fontName="Helvetica-Bold",
                                   textColor=c_h1, leading=(fs["title"]+8)*1.25, alignment=TA_CENTER)
        cover_s_s = ParagraphStyle("CS", fontSize=fs["body"]+1, fontName="Helvetica",
                                   textColor=c_muted, alignment=TA_CENTER, spaceAfter=6)
        story += [
            Spacer(1, 5*cm),
            _p(doc_title, cover_t_s),
            Spacer(1, 6*mm),
            HRFlowable(width="50%", thickness=3, color=c_header, hAlign="CENTER"),
            Spacer(1, 4*mm),
            _p("Erstellt mit Xera AI", cover_s_s),
            PageBreak(),
        ]

    # — Optional table of contents —
    if toc:
        toc_title_s = ParagraphStyle("TOCT", fontSize=fs["title"]-4, fontName="Helvetica-Bold",
                                     textColor=c_ink, leading=(fs["title"]-4)*1.15, spaceAfter=3*sm)
        toc_obj = TableOfContents()
        toc_obj.levelStyles = [
            ParagraphStyle("TOC0", fontSize=fs["h2"], fontName="Helvetica-Bold",
                          textColor=c_ink, leftIndent=0, firstLineIndent=0,
                          spaceBefore=8, leading=fs["h2"]*1.4),
            ParagraphStyle("TOC1", fontSize=fs["body"], fontName="Helvetica",
                          textColor=c_muted, leftIndent=16, firstLineIndent=0,
                          spaceBefore=2, leading=fs["body"]*1.4),
        ]
        story += [
            _p("Inhaltsverzeichnis", toc_title_s),
            HRFlowable(width="100%", thickness=2, color=c_header, spaceAfter=10*sm),
            toc_obj,
            PageBreak(),
        ]

    # — Main parse loop —
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # ── Code block ──
        if line.startswith("```"):
            lang = line[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].rstrip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                code_text = _wrap_code("\n".join(code_lines))
                col_w = doc.width - 2*PAD

                # Language label header
                header_rows = []
                if lang:
                    lang_s = ParagraphStyle("LNG", fontSize=7.5, fontName="Helvetica-Bold",
                                            textColor=colors.HexColor(theme["header_text"]))
                    lang_cell = _p(lang.upper(), lang_s)
                    lang_row = Table([[lang_cell]], colWidths=[col_w])
                    lang_row.setStyle(TableStyle([
                        ("BACKGROUND",    (0,0), (-1,-1), c_header),
                        ("LEFTPADDING",   (0,0), (-1,-1), PAD),
                        ("RIGHTPADDING",  (0,0), (-1,-1), PAD),
                        ("TOPPADDING",    (0,0), (-1,-1), 4),
                        ("BOTTOMPADDING", (0,0), (-1,-1), 4),
                    ]))
                    story.append(Spacer(1, 4*sm))
                    story.append(lang_row)
                else:
                    story.append(Spacer(1, 4*sm))

                code_tbl = Table([[Preformatted(code_text, code_s)]], colWidths=[col_w])
                code_tbl.setStyle(TableStyle([
                    ("BACKGROUND",    (0,0), (-1,-1), c_codebg),
                    ("BOX",           (0,0), (-1,-1), 0.75, c_codebrd),
                    ("LEFTPADDING",   (0,0), (-1,-1), PAD),
                    ("RIGHTPADDING",  (0,0), (-1,-1), PAD),
                    ("TOPPADDING",    (0,0), (-1,-1), 8),
                    ("BOTTOMPADDING", (0,0), (-1,-1), 8),
                ]))
                story.append(code_tbl)
                story.append(Spacer(1, 6*sm))
            i += 1
            continue

        # ── Callout / Blockquote  >  ──
        if line.startswith("> ") or line.strip() == ">":
            quote_lines = []
            while i < len(lines) and (lines[i].startswith("> ") or lines[i].strip() == ">"):
                q = lines[i][2:] if lines[i].startswith("> ") else ""
                quote_lines.append(q)
                i += 1
            # Parse [!type] or [!type] Title
            ctype, ctitle = None, None
            if quote_lines and re.match(r'^\[!(\w+)\]', quote_lines[0]):
                m = re.match(r'^\[!(\w+)\]\s*(.*)', quote_lines[0])
                ctype = m.group(1).lower()
                ctitle = m.group(2).strip() or None
                quote_lines = quote_lines[1:]
            box = _make_callout(quote_lines, ctype, ctitle,
                                theme["accent"], theme["accent_light"],
                                doc.width, fs["body"], lm)
            if box:
                story.append(Spacer(1, 5*sm))
                story.append(box)
                story.append(Spacer(1, 5*sm))
            continue

        # ── Markdown table ──
        if "|" in line and line.strip().startswith("|"):
            tbl_lines = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                if not re.match(r'^[\s|:\-]+$', lines[i]):
                    tbl_lines.append(lines[i])
                i += 1
            if tbl_lines:
                rows = [[c.strip() for c in tl.strip().strip("|").split("|")] for tl in tbl_lines]
                col_count = max(len(r) for r in rows)
                rows = [r + [""] * (col_count - len(r)) for r in rows]
                col_w = doc.width / col_count
                tbl_data = [
                    [_p(cell, ParagraphStyle("TC",
                        fontSize=fs["body"]-1,
                        fontName="Helvetica-Bold" if ri == 0 else "Helvetica",
                        textColor=white if ri == 0 else c_body,
                        leading=(fs["body"]-1)*1.35,
                        alignment=TA_CENTER if ri == 0 else TA_LEFT,
                        wordWrap="LTR"))
                     for cell in row]
                    for ri, row in enumerate(rows)
                ]
                tbl = Table(tbl_data, colWidths=[col_w]*col_count, repeatRows=1)
                tbl.setStyle(TableStyle([
                    ("BACKGROUND",     (0,0), (-1, 0),  c_tblhd),
                    ("ROWBACKGROUNDS", (0,1), (-1,-1),  [white, c_tblalt]),
                    ("GRID",           (0,0), (-1,-1),  0.4, c_hr),
                    ("BOX",            (0,0), (-1,-1),  1,   c_tblhd),
                    ("LEFTPADDING",    (0,0), (-1,-1),  8),
                    ("RIGHTPADDING",   (0,0), (-1,-1),  8),
                    ("TOPPADDING",     (0,0), (-1,-1),  6),
                    ("BOTTOMPADDING",  (0,0), (-1,-1),  6),
                    ("VALIGN",         (0,0), (-1,-1),  "MIDDLE"),
                ]))
                story.append(Spacer(1, 6*sm))
                story.append(tbl)
                story.append(Spacer(1, 8*sm))
            continue

        # ── Horizontal rule ──
        if re.match(r'^[-*_]{3,}\s*$', line):
            story.append(Spacer(1, 6*sm))
            story.append(HRFlowable(width="100%", thickness=1, color=c_hr, spaceAfter=6*sm))
            i += 1
            continue

        # ── H1 title ──
        if line.startswith("# ") and not line.startswith("## "):
            story.append(Spacer(1, 6*sm))
            story.append(_p(line[2:].strip(), title_s))
            story.append(HRFlowable(width="100%", thickness=2.5, color=c_header, spaceAfter=8*sm))

        # ── H2 section header (colored strip) ──
        elif line.startswith("## ") and not line.startswith("### "):
            story.append(Spacer(1, 16*sm))
            sec_header = _make_section_header(line[3:].strip(), theme["h1"], doc.width, fs["h1"], lm)
            if toc:
                sec_header._toc_entry = (0, line[3:].strip())
            story.append(sec_header)
            story.append(Spacer(1, 7*sm))

        # ── H3 ──
        elif line.startswith("### ") and not line.startswith("#### "):
            story.append(Spacer(1, 10*sm))
            h3_para = _p(line[4:].strip(), h2_s)
            if toc:
                h3_para._toc_entry = (1, line[4:].strip())
            story.append(h3_para)
            story.append(HRFlowable(width="40%", thickness=0.5, color=c_hr, spaceAfter=3*sm))

        # ── H4 ──
        elif line.startswith("#### "):
            story.append(_p(line[5:].strip(), h3_s))

        # ── Sub-bullet (2+ spaces + -) ──
        elif re.match(r'^  +[-*+] ', line):
            text = re.sub(r'^  +[-*+] ', '', line)
            story.append(_p("– " + text.strip(), blt2_s))

        # ── Bullet ──
        elif re.match(r'^[-*+] ', line):
            story.append(_p("• " + line[2:].strip(), blt_s))

        # ── Numbered list ──
        elif re.match(r'^\d+\.\s', line):
            m = re.match(r'^(\d+)\.\s+(.*)', line)
            if m:
                story.append(_p(f"{m.group(1)}.  {m.group(2).strip()}", num_s))

        # ── Empty line ──
        elif not line.strip():
            story.append(Spacer(1, 4*sm))

        # ── Normal paragraph ──
        else:
            story.append(_p(line, body_s))

        i += 1

    if toc:
        doc.multiBuild(story)
    else:
        doc.build(
            story,
            onFirstPage=lambda c, d: _draw_header(c, d, first=True),
            onLaterPages=lambda c, d: _draw_header(c, d, first=False),
        )


# ─── DOCX ─────────────────────────────────────────────────────────────────────

def _create_docx(content: str, path: Path, theme: dict):
    from docx import Document
    from docx.shared import RGBColor, Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    def hex2rgb(h):
        h = h.lstrip("#")
        return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

    accent = hex2rgb(theme["accent"])
    h1_rgb = hex2rgb(theme["h1"])

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3);  sec.right_margin  = Cm(3)

    for line in content.split("\n"):
        line = line.rstrip()
        if not line:
            doc.add_paragraph()
        elif line.startswith("# ") and not line.startswith("## "):
            p = doc.add_heading(line[2:], level=1)
            for run in p.runs: run.font.color.rgb = RGBColor(*accent)
        elif line.startswith("## ") and not line.startswith("### "):
            p = doc.add_heading(line[3:], level=2)
            for run in p.runs: run.font.color.rgb = RGBColor(*h1_rgb)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif re.match(r'^  +[-*+] ', line):
            text = re.sub(r'^  +[-*+] ', '', line)
            doc.add_paragraph(text, style="List Bullet 2")
        elif re.match(r'^[-*+] ', line):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r'^\d+\.\s', line):
            m = re.match(r'^\d+\.\s+(.*)', line)
            if m: doc.add_paragraph(m.group(1), style="List Number")
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:])
            p.style = doc.styles["Quote"] if "Quote" in [s.name for s in doc.styles] else p.style
        else:
            doc.add_paragraph(line)

    doc.save(str(path))


# ─── XLSX ─────────────────────────────────────────────────────────────────────

def _create_xlsx(content: str, path: Path, theme: dict):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    def hfill(h):
        return PatternFill(start_color=h.lstrip("#"), end_color=h.lstrip("#"), fill_type="solid")

    wb = openpyxl.Workbook()
    ws = wb.active
    thin = Side(style="thin", color=theme["code_border"].lstrip("#"))
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    lines = [l for l in content.split("\n") if l.strip()]
    for ri, line in enumerate(lines, start=1):
        cells = [c.strip() for c in (line.split("|") if "|" in line else line.split(","))]
        for ci, val in enumerate(cells, start=1):
            cell = ws.cell(row=ri, column=ci, value=val)
            cell.border = border
            if ri == 1:
                cell.font = Font(bold=True, color="FFFFFF", size=11)
                cell.fill = hfill(theme["tbl_head"])
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            else:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.fill = hfill(theme["tbl_alt"]) if ri % 2 == 0 else PatternFill()

    for col in ws.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=8)
        ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 4, 55)
    ws.freeze_panes = "A2"

    wb.save(str(path))
